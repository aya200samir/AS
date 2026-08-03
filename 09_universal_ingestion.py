"""Memory-aware, session-scoped multimodal ingestion for AS Intelligence Studio."""

from contextlib import contextmanager
from io import BytesIO
from pathlib import Path
import base64
import hashlib
import os
import re
import tempfile

import fitz
import numpy as np
import pandas as pd
from docx import Document
from openpyxl import load_workbook
from PIL import Image
from sklearn.metrics.pairwise import cosine_similarity


SUPPORTED_EXTENSIONS = {
    "pdf", "docx", "xlsx", "xls", "csv", "png", "jpg", "jpeg", "webp"
}
MAX_FILE_BYTES = 100 * 1024 * 1024
MAX_TOTAL_BYTES = 200 * 1024 * 1024
MAX_IMAGE_BYTES = 25 * 1024 * 1024
MAX_FILES = 6
MAX_VISUALS_PER_FILE = 6
MAX_CHUNKS_PER_FILE = 12_000
MAX_TOTAL_CHUNKS = 30_000
COPY_BLOCK_BYTES = 1024 * 1024
EMBEDDING_BATCH_SIZE = 32


def _clean(text):
    text = str(text or "").replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _chunks(text, source_name, location, kind="text", target_words=210, overlap=35):
    words = _clean(text).split()
    if not words:
        return []
    rows, start, part = [], 0, 1
    while start < len(words):
        end = min(start + target_words, len(words))
        payload = " ".join(words[start:end])
        identity = f"{source_name}|{location}|{part}|{payload}".encode("utf-8")
        rows.append({
            "chunk_id": hashlib.sha256(identity).hexdigest()[:24],
            "source_name": source_name,
            "location": location,
            "content_kind": kind,
            "text": payload,
            "word_count": len(payload.split()),
        })
        if end == len(words):
            break
        start = max(start + 1, end - overlap)
        part += 1
    return rows


def _image_record(image_bytes, source_name, location, mime_type):
    digest = hashlib.sha256(image_bytes).hexdigest()[:24]
    return {
        "visual_id": digest,
        "source_name": source_name,
        "location": location,
        "mime_type": mime_type,
        "bytes": image_bytes,
    }


def _reported_size(uploaded):
    size = getattr(uploaded, "size", None)
    if size is not None:
        return int(size)
    try:
        position = uploaded.tell()
        uploaded.seek(0, os.SEEK_END)
        size = uploaded.tell()
        uploaded.seek(position)
        return int(size)
    except Exception:
        return 0


@contextmanager
def _temporary_upload(uploaded, suffix, byte_limit):
    """Copy an upload to disk in bounded blocks and always remove it afterwards."""
    path = None
    try:
        uploaded.seek(0)
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as target:
            path = Path(target.name)
            total = 0
            while True:
                block = uploaded.read(COPY_BLOCK_BYTES)
                if not block:
                    break
                total += len(block)
                if total > byte_limit:
                    raise ValueError(
                        f"File exceeds the {byte_limit // (1024 * 1024)} MB limit"
                    )
                target.write(block)
        yield path, total
    finally:
        try:
            uploaded.seek(0)
        except Exception:
            pass
        if path and path.exists():
            path.unlink(missing_ok=True)


def _append_bounded(target, additions, limit=MAX_CHUNKS_PER_FILE):
    remaining = max(0, limit - len(target))
    if remaining:
        target.extend(additions[:remaining])
    return len(target) >= limit


def _extract_pdf(path, name, progress_callback=None):
    rows, visuals = [], []
    document = fitz.open(path)
    total_pages = max(document.page_count, 1)
    try:
        for page_index, page in enumerate(document, start=1):
            if progress_callback:
                progress_callback(page_index, total_pages, f"Reading {name}: page {page_index}/{total_pages}")
            reached_limit = _append_bounded(
                rows, _chunks(page.get_text("text"), name, f"Page {page_index}")
            )
            if len(visuals) < MAX_VISUALS_PER_FILE:
                needed = MAX_VISUALS_PER_FILE - len(visuals)
                for image_index, image in enumerate(page.get_images(full=True)[:needed], start=1):
                    try:
                        payload = document.extract_image(image[0])
                        image_bytes = payload["image"]
                        if len(image_bytes) <= MAX_IMAGE_BYTES:
                            visuals.append(_image_record(
                                image_bytes,
                                name,
                                f"Page {page_index}, image {image_index}",
                                f"image/{payload.get('ext', 'png')}",
                            ))
                    except Exception:
                        continue
            if reached_limit:
                break
    finally:
        document.close()
    return rows, visuals


def _extract_docx(path, name, progress_callback=None):
    rows, visuals = [], []
    document = Document(path)
    total_items = max(len(document.paragraphs) + len(document.tables), 1)
    completed = 0
    for index, paragraph in enumerate(document.paragraphs, start=1):
        completed += 1
        if progress_callback and completed % 50 == 0:
            progress_callback(completed, total_items, f"Reading {name}: paragraph {index}")
        if _append_bounded(rows, _chunks(paragraph.text, name, f"Paragraph {index}")):
            break
    if len(rows) < MAX_CHUNKS_PER_FILE:
        for table_index, table in enumerate(document.tables, start=1):
            completed += 1
            table_text = "\n".join(
                " | ".join(cell.text for cell in row.cells) for row in table.rows
            )
            if _append_bounded(
                rows, _chunks(table_text, name, f"Table {table_index}", "table")
            ):
                break
    for index, relation in enumerate(document.part.rels.values(), start=1):
        if "image" in relation.reltype and len(visuals) < MAX_VISUALS_PER_FILE:
            blob = relation.target_part.blob
            if len(blob) <= MAX_IMAGE_BYTES:
                visuals.append(
                    _image_record(blob, name, f"Embedded image {index}", "image/png")
                )
    return rows, visuals


def _extract_csv(path, name, progress_callback=None):
    rows, row_count, columns = [], 0, []
    for part, frame in enumerate(
        pd.read_csv(path, chunksize=5_000, low_memory=False), start=1
    ):
        frame = frame.dropna(how="all").dropna(axis=1, how="all")
        if frame.empty:
            continue
        if not columns:
            columns = list(map(str, frame.columns))
        row_count += len(frame)
        if progress_callback:
            progress_callback(part, part + 1, f"Reading {name}: {row_count:,} rows")
        if _append_bounded(
            rows,
            _chunks(frame.astype(str).to_csv(index=False), name, f"CSV rows through {row_count:,}", "table"),
        ):
            break
    summary = f"CSV. Rows read: {row_count}. Columns: {', '.join(columns)}."
    _append_bounded(rows, _chunks(summary, name, "CSV summary", "table_summary"))
    return rows, []


def _extract_xlsx(path, name, progress_callback=None):
    rows = []
    workbook = load_workbook(path, read_only=True, data_only=True)
    try:
        for sheet_index, sheet in enumerate(workbook.worksheets, start=1):
            buffer, row_count = [], 0
            for values in sheet.iter_rows(values_only=True):
                if not any(value is not None for value in values):
                    continue
                buffer.append(" | ".join("" if value is None else str(value) for value in values))
                row_count += 1
                if len(buffer) >= 250:
                    if _append_bounded(
                        rows,
                        _chunks("\n".join(buffer), name, f"Sheet: {sheet.title}, rows through {row_count}", "table"),
                    ):
                        break
                    buffer.clear()
                if progress_callback and row_count % 2_000 == 0:
                    progress_callback(sheet_index, len(workbook.worksheets), f"Reading {name}: {sheet.title}, {row_count:,} rows")
            if buffer and len(rows) < MAX_CHUNKS_PER_FILE:
                _append_bounded(
                    rows,
                    _chunks("\n".join(buffer), name, f"Sheet: {sheet.title}, final rows", "table"),
                )
            _append_bounded(
                rows,
                _chunks(
                    f"Sheet {sheet.title}. Rows read: {row_count}.",
                    name,
                    f"Sheet: {sheet.title} summary",
                    "table_summary",
                ),
            )
            if len(rows) >= MAX_CHUNKS_PER_FILE:
                break
    finally:
        workbook.close()
    return rows, []


def _extract_xls(path, name, progress_callback=None):
    rows = []
    book = pd.ExcelFile(path)
    for sheet_index, sheet_name in enumerate(book.sheet_names, start=1):
        if progress_callback:
            progress_callback(sheet_index, len(book.sheet_names), f"Reading {name}: {sheet_name}")
        frame = pd.read_excel(book, sheet_name=sheet_name)
        frame = frame.dropna(how="all").dropna(axis=1, how="all")
        if frame.empty:
            continue
        if _append_bounded(
            rows,
            _chunks(frame.astype(str).to_csv(index=False), name, f"Sheet: {sheet_name}", "table"),
        ):
            break
    return rows, []


def ingest_uploaded_files(uploaded_files, progress_callback=None):
    if len(uploaded_files) > MAX_FILES:
        raise ValueError(f"Upload at most {MAX_FILES} files per workspace.")
    reported_total = sum(_reported_size(uploaded) for uploaded in uploaded_files)
    if reported_total > MAX_TOTAL_BYTES:
        raise ValueError("The combined workspace exceeds the 200 MB session limit.")

    rows, visuals, report = [], [], []
    total_files = max(len(uploaded_files), 1)
    for file_index, uploaded in enumerate(uploaded_files, start=1):
        name = Path(uploaded.name).name
        extension = name.rsplit(".", 1)[-1].lower() if "." in name else ""
        size = _reported_size(uploaded)
        if progress_callback:
            progress_callback(file_index - 1, total_files, f"Preparing {name}")
        if extension not in SUPPORTED_EXTENSIONS:
            report.append({"file": name, "status": "Rejected", "detail": "Unsupported file type"})
            continue
        byte_limit = MAX_IMAGE_BYTES if extension in {"png", "jpg", "jpeg", "webp"} else MAX_FILE_BYTES
        if size and size > byte_limit:
            report.append({
                "file": name,
                "status": "Rejected",
                "detail": f"File exceeds the {byte_limit // (1024 * 1024)} MB limit",
            })
            continue
        try:
            with _temporary_upload(uploaded, f".{extension}", byte_limit) as (path, actual_size):
                nested_progress = None
                if progress_callback:
                    nested_progress = lambda done, total, message: progress_callback(
                        min(file_index - 1 + (done / max(total, 1)), total_files),
                        total_files,
                        message,
                    )
                if extension == "pdf":
                    file_rows, file_visuals = _extract_pdf(path, name, nested_progress)
                elif extension == "docx":
                    file_rows, file_visuals = _extract_docx(path, name, nested_progress)
                elif extension == "csv":
                    file_rows, file_visuals = _extract_csv(path, name, nested_progress)
                elif extension == "xlsx":
                    file_rows, file_visuals = _extract_xlsx(path, name, nested_progress)
                elif extension == "xls":
                    file_rows, file_visuals = _extract_xls(path, name, nested_progress)
                else:
                    with Image.open(path) as image:
                        image.verify()
                    image_bytes = path.read_bytes()
                    file_rows = []
                    file_visuals = [_image_record(
                        image_bytes, name, "Standalone image",
                        uploaded.type or f"image/{extension}",
                    )]

            remaining = max(0, MAX_TOTAL_CHUNKS - len(rows))
            accepted_rows = file_rows[:remaining]
            rows.extend(accepted_rows)
            visuals.extend(file_visuals)
            truncated = len(accepted_rows) < len(file_rows) or len(file_rows) >= MAX_CHUNKS_PER_FILE
            if accepted_rows:
                detail = f"{len(accepted_rows):,} text chunks, {len(file_visuals)} visuals"
                if truncated:
                    detail += " (safe processing limit reached)"
                status = "Ready"
            elif file_visuals:
                detail = f"0 text chunks, {len(file_visuals)} visuals; use Visual Intelligence"
                status = "Visual only"
            elif extension == "pdf":
                detail = "No searchable text found; this may be a scanned PDF requiring OCR"
                status = "Needs OCR"
            else:
                detail = "No searchable content was extracted"
                status = "Empty"
            report.append({
                "file": name,
                "size_mb": round(actual_size / (1024 * 1024), 1),
                "status": status,
                "detail": detail,
            })
            if len(rows) >= MAX_TOTAL_CHUNKS:
                break
        except Exception as error:
            report.append({
                "file": name,
                "size_mb": round(size / (1024 * 1024), 1) if size else None,
                "status": "Failed",
                "detail": str(error)[:220],
            })
        if progress_callback:
            progress_callback(file_index, total_files, f"Finished {name}")
    return pd.DataFrame(rows), visuals, pd.DataFrame(report)


def embed_workspace(chunks_df, embedding_model, progress_callback=None):
    if chunks_df.empty:
        return np.empty((0, 384), dtype=np.float32)
    texts = chunks_df["text"].tolist()
    batches = []
    total_batches = (len(texts) + EMBEDDING_BATCH_SIZE - 1) // EMBEDDING_BATCH_SIZE
    for batch_index, start in enumerate(range(0, len(texts), EMBEDDING_BATCH_SIZE), start=1):
        if progress_callback:
            progress_callback(batch_index - 1, total_batches, f"Embedding batch {batch_index}/{total_batches}")
        batches.append(embedding_model.encode(
            texts[start:start + EMBEDDING_BATCH_SIZE],
            normalize_embeddings=True,
            show_progress_bar=False,
            batch_size=EMBEDDING_BATCH_SIZE,
        ))
    if progress_callback:
        progress_callback(total_batches, total_batches, "Embeddings ready")
    return np.vstack(batches).astype(np.float32, copy=False)


def retrieve_workspace(question, chunks_df, embeddings, embedding_model, top_k=7):
    if chunks_df.empty or not len(embeddings):
        return chunks_df.head(0).copy()
    query = embedding_model.encode(
        [question], normalize_embeddings=True, show_progress_bar=False
    )
    scores = cosine_similarity(query, embeddings)[0]
    take = np.argsort(scores)[::-1][:min(top_k, len(scores))]
    result = chunks_df.iloc[take].copy()
    result["score"] = scores[take]
    return result.reset_index(drop=True)


def build_evidence_context(results):
    blocks = []
    for index, row in results.reset_index(drop=True).iterrows():
        blocks.append(
            f"[Evidence {index + 1}]\nFile: {row['source_name']}\n"
            f"Location: {row['location']}\nType: {row['content_kind']}\nText: {row['text']}"
        )
    return "\n\n".join(blocks)


def image_data_uri(visual):
    encoded = base64.b64encode(visual["bytes"]).decode("ascii")
    return f"data:{visual['mime_type']};base64,{encoded}"
